import docx
import pandas as pd
import numpy as np
import re
from tinydb import TinyDB
db=TinyDB('sampleDoc2.json')
def readSampleInputDoc1():
	filename='SampleInputDoc1-FAQs.docx'
	try:
		doc = docx.Document(filename)
		QnA=[]
		completeAns=[]
		status = False
		for para in doc.paragraphs:
			paraText=para.text
			if status:
				completeAns.append(paraText)
			if paraText.endswith('?'):
				
				if status and completeAns:
					ans=''.join(completeAns[:-1])
					QnA.append({'Question':ques,'Answer':ans})
					completeAns=[]
				ques=paraText
				status=True
			else:
				for run in para.runs:
					if run.bold:
						if status and completeAns:
							ans=''.join(completeAns[:-1])
							QnA.append({'Question':ques,'Answer':ans})
							completeAns=[]
						ques=paraText
						status=True
						break
		ans=''.join(completeAns)
		QnA.append({'Question':ques,'Answer':ans})
		return QnA
	except IOError:
		print('File does not exist to appear')
def readSampleInputDoc3(source=False):
    filename="SampleInputDoc3-Hardware Problems.docx"
    symptomFlag= False
    diagnosisFlag=False
    symptomContent=[]
    diagnosisContent=[]
    QnA=[]
    try:
        doc=docx.Document(filename)
        mannualDict={}
        for para in doc.paragraphs:
            paraText=para.text
            if symptomFlag:
                symptomContent.append(paraText)
            elif diagnosisFlag:
                diagnosisContent.append(paraText)
            if 'troubleshooting' in paraText.lower():
                if diagnosisFlag:
                    ques=topic +' '+''.join(symptomContent[:-1])
                    ans='\n'.join(diagnosisContent[:-1])
                    QnA.append({'Question':ques,'Answer':ans})
                    myDict={'symptom':'\n'.join(symptomContent[:-1]),'diagnosis':'\n'.join(diagnosisContent[:-1])}
                    if topic in mannualDict:
                        mannualDict[topic].append(myDict)
                    else:
                        mannualDict[topic]=[myDict]
                    symptomContent = []
                    diagnosisContent = []
                    diagnosisFlag=False
                topic = paraText.replace('/',' ')
            elif 'symptom' in paraText.lower():
                if diagnosisFlag:
                    ques=topic +' '+''.join(symptomContent[:-1])
                    ans='\n'.join(diagnosisContent[:-1])
                    QnA.append({'Question':ques,'Answer':ans})
                    myDict={'symptom':'\n'.join(symptomContent[:-1]),'diagnosis':'\n'.join(diagnosisContent[:-1])}
                    if topic in mannualDict:
                        mannualDict[topic].append(myDict)
                    else:
                        mannualDict[topic]=[myDict]
                    symptomContent = []
                    diagnosisContent = []
                symptomFlag=True
                diagnosisFlag=False
            elif 'diagnosis' in paraText.lower():
                diagnosisFlag=True
                symptomFlag=False
        if source:
            return mannualDict
        #return mannualDict
        return QnA
                
    except IOError:
        print('File does not exist to appear.')
def readSampleInput():
    filename="SampleInput.xlsx"
    try:
        xl=pd.ExcelFile(filename)
        sheetNames=xl.sheet_names
        df=xl.parse(sheetNames[0])
        response=df.groupby('TicketNumber')['TicketNumber'].filter(lambda x: len(x) == 1)
        
        df1=pd.DataFrame(response,columns=['TicketNumber'])
        new_df=df1.join(df,how='inner',rsuffix='second')
        new_df=new_df.dropna(subset=['Title'])
        QnA=[]
        QnADict={}
        for index , row in new_df.iterrows():
            TicketNumber=row['TicketNumbersecond']
            Title=row['Title']
            Resolution=row['Resolution']
            HelpTopic=row['HelpTopic']
            if Title.lower()=='status changed' or Title.lower()=='collaborators added by end user' :
                continue
            if '<br />' in Resolution.lower():
                resp=re.findall(r'<br /><br />(.+?)<br />',Resolution.lower()+'<br />')
                Resolution=resp[0]
            if Resolution in QnADict:
                QnADict[Resolution]+=' '+Title
            else:
                QnADict[Resolution]=Title
        for Resolution,Title in QnADict.items():
            QnA.append({'Question':Title,'Answer':Resolution})
        return QnA

        

    except IOError:
        print("File does not exist to appear.")
def readSampleInputDoc2():
    filename="SampleInputDoc2-.docx"
    try:
        doc=docx.Document(filename)
        QnA=[]
        status=False
        answer=""
        for para in doc.paragraphs:
            paraText=para.text
            num=0
            for run in para.runs:
                num+=1
            if num==1 and len(paraText.split()) < 20:
                if status:
                    QnA.append({'Question':ques,'Answer':answer})
                    answer=""
                ques=paraText
                status=True
            else:
                answer+=paraText+"\n"


        QnA.append({'Question':ques,'Answer':answer})
        return QnA



    except IOError:
        print("File does not exist to appear.")
def dataForSampleInputDoc2FromDB():
    resp=db.all()
    QnA=[]
    for item in resp:
        ques=item['root']
        answer=""
        for x in item['topic']:
            answer+="<strong>"+x['problem']+"</strong><br>"+x['solution']+"<br/>"
        QnA.append({'Answer':answer,'Question':ques})
    return QnA

#dataForSampleInputDoc2FromDB()
#readSampleInput()
#path='SampleInputDoc1-FAQs.docx'
#response=readSampleInputDoc3()
#readSampleDoc2()


